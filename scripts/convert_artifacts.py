#!/usr/bin/env python3
"""
Raw Artifact Converter
Converts various file types to Markdown for Copilot processing.

Usage: python convert_artifacts.py [--input INPUT_DIR] [--output OUTPUT_DIR]

Supported conversions:
- PDF (.pdf) → Markdown (.md)
- Word (.docx) → Markdown (.md)
- Text files with encoding fixes
- CSV → Markdown tables (preserved)
- Images → described in .md (requires manual description)
"""

import os
import sys
import argparse
import subprocess
from pathlib import Path
from datetime import datetime

# Optional dependencies - will work without them
pdftotext = None
Document = None

try:
    import pdftotext
except ImportError:
    pass

try:
    import markdown
    from docx import Document
except ImportError:
    pass

# ANSI colors
GREEN = '\033[92m'
YELLOW = '\033[93m'
RED = '\033[91m'
RESET = '\033[0m'

def log_info(msg):
    print(f"{GREEN}[INFO]{RESET} {msg}")

def log_warn(msg):
    print(f"{YELLOW}[WARN]{RESET} {msg}")

def log_error(msg):
    print(f"{RED}[ERROR]{RESET} {msg}")

class ArtifactConverter:
    def __init__(self, input_dir, output_dir):
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.converted_files = []
        self.skipped_files = []
        self.errors = []
        
    def convert_pdf(self, pdf_path, output_path):
        """Convert PDF to Markdown using pdftotext."""
        if pdftotext is None:
            self._create_fallback_conversion(pdf_path, output_path, "PDF")
            return
            
        try:
            with open(pdf_path, 'rb') as f:
                pdf = pdftotext.PDF(f)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"# Converted from: {pdf_path.name}\n")
                f.write(f"# Converted: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")
                
                for i, page in enumerate(pdf):
                    f.write(f"## Page {i+1}\n\n")
                    f.write(page)
                    f.write("\n\n")
                    
            log_info(f"Converted PDF: {pdf_path.name}")
            self.converted_files.append(str(output_path))
        except Exception as e:
            self._create_fallback_conversion(pdf_path, output_path, f"PDF - Error: {e}")

    def convert_docx(self, docx_path, output_path):
        """Convert Word document to Markdown."""
        if Document is None:
            self._create_fallback_conversion(docx_path, output_path, "DOCX (python-docx not installed)")
            return
            
        try:
            doc = Document(docx_path)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"# Converted from: {docx_path.name}\n")
                f.write(f"# Converted: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        # Simple heading detection
                        if para.style.name.startswith('Heading'):
                            level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
                            f.write(f"{'#' * int(level)} {para.text}\n\n")
                        else:
                            f.write(f"{para.text}\n\n")
                
                # Handle tables
                for table in doc.tables:
                    f.write("\n| ")
                    for cell in table.rows[0].cells:
                        f.write(f"{cell.text} |")
                    f.write("\n| ")
                    for _ in table.rows[0].cells:
                        f.write("---|")
                    f.write("\n")
                    for row in table.rows[1:]:
                        f.write("| ")
                        for cell in row.cells:
                            f.write(f"{cell.text} |")
                        f.write("\n")
                        
            log_info(f"Converted DOCX: {docx_path.name}")
            self.converted_files.append(str(output_path))
        except Exception as e:
            self._create_fallback_conversion(docx_path, output_path, f"DOCX - Error: {e}")

    def convert_txt(self, txt_path, output_path):
        """Convert text file with encoding fixes."""
        try:
            # Try different encodings
            for enc in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(txt_path, 'r', encoding=enc) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                with open(txt_path, 'rb') as f:
                    content = f.read().decode('utf-8', errors='replace')
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"# Converted from: {txt_path.name}\n\n")
                f.write(content)
                
            log_info(f"Converted TXT: {txt_path.name}")
            self.converted_files.append(str(output_path))
        except Exception as e:
            log_error(f"Failed to convert {txt_path.name}: {e}")
            self.errors.append((str(txt_path), str(e)))

    def create_image_markdown(self, img_path, output_path):
        """Create markdown reference for images."""
        relative_path = f"images/{img_path.name}"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"# Image: {img_path.name}\n\n")
            f.write(f"![{img_path.name}](../raw-artifacts/images/{img_path.name})\n\n")
            f.write("## Description\n")
            f.write("_TODO: Describe what this diagram/screenshot shows_\n\n")
            f.write("## Key Elements\n")
            f.write("- _TODO: List key elements visible in image_\n")
            
        log_info(f"Created image reference: {img_path.name}")
        self.converted_files.append(str(output_path))

    def _create_fallback_conversion(self, source_path, output_path, type_name):
        """Create a placeholder when conversion isn't possible."""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"# Conversion Required: {source_path.name}\n\n")
            f.write(f"**Original file:** {source_path.name}\n")
            f.write(f"**File type:** {type_name}\n\n")
            f.write("## Manual Conversion Needed\n\n")
            f.write("Please copy the content from this file into markdown format:\n\n")
            f.write("1. Open the original file\n")
            f.write("2. Copy text content\n")
            f.write("3. Paste into this file with proper markdown formatting\n")
            f.write("4. Delete this notice\n\n")
            
        log_warn(f"Manual conversion needed: {source_path.name}")
        self.skipped_files.append(str(output_path))

    def copy_markdown(self, md_path, output_path):
        """Copy markdown files as-is."""
        with open(md_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        log_info(f"Copied: {md_path.name}")
        self.converted_files.append(str(output_path))

    def copy_csv(self, csv_path, output_path):
        """Copy CSV files (Copilot can read them)."""
        with open(csv_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("# CSV Data\n\n")
            f.write("```csv\n")
            f.write(content)
            f.write("\n```\n")
        log_info(f"Copied CSV: {csv_path.name}")
        self.converted_files.append(str(output_path))

    def process(self):
        """Process all files in input directory."""
        log_info(f"Scanning: {self.input_dir}")
        
        # Ensure output directory exists
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Create images directory
        images_dir = self.input_dir / "images"
        images_dir.mkdir(exist_ok=True)
        
        # File type handlers
        handlers = {
            '.pdf': self.convert_pdf,
            '.docx': self.convert_docx,
            '.txt': self.convert_txt,
            '.md': self.copy_markdown,
            '.csv': self.copy_csv,
        }
        
        image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp'}
        
        # Process all files
        for file_path in self.input_dir.rglob('*'):
            if not file_path.is_file():
                continue
            if file_path.name.startswith('.'):
                continue
            if 'images' in file_path.parts:
                continue
                
            relative_path = file_path.relative_to(self.input_dir)
            output_path = self.output_dir / relative_path
            
            # Create parent directories
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            suffix = file_path.suffix.lower()
            
            if suffix in handlers:
                handlers[suffix](file_path, output_path.with_suffix('.md'))
            elif suffix in image_extensions:
                self.create_image_markdown(file_path, output_path.with_suffix('.md'))
            else:
                log_warn(f"Unknown file type: {file_path.name}")
                
        # Summary
        print(f"\n{'='*50}")
        log_info(f"Conversion complete!")
        log_info(f"  Converted: {len(self.converted_files)} files")
        log_info(f"  Skipped/Needs manual: {len(self.skipped_files)} files")
        log_info(f"  Errors: {len(self.errors)} files")
        print(f"{'='*50}\n")
        
        if self.skipped_files:
            log_warn("Files needing manual conversion:")
            for f in self.skipped_files:
                print(f"  - {f}")
                
        return len(self.errors) == 0

def main():
    parser = argparse.ArgumentParser(description='Convert raw artifacts to Markdown')
    parser.add_argument('--input', default='raw-artifacts', help='Input directory')
    parser.add_argument('--output', default='raw-artifacts/converted', help='Output directory')
    args = parser.parse_args()
    
    converter = ArtifactConverter(args.input, args.output)
    success = converter.process()
    sys.exit(0 if success else 1)

if __name__ == '__main__':
    main()
